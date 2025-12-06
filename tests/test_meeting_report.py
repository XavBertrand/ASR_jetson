# tests/test_meeting_report.py
from __future__ import annotations
import json
from pathlib import Path

import pytest

# On importe la fonction orchestratrice
# (le chemin ci-dessous suit la structure que je t'ai proposée)
from src.asr_jetson.postprocessing.meeting_report import generate_meeting_report

# On importera les modules à monkeypatcher via leurs chemins réels
import src.asr_jetson.postprocessing.mistral_client as mistral_client_mod
import src.asr_jetson.postprocessing.meeting_report as meeting_report_mod

try:
    from docx import Document
except Exception as exc:  # pragma: no cover - dépend de l'environnement CI
    Document = None  # type: ignore
    _DOCX_IMPORT_ERROR = exc
else:
    _DOCX_IMPORT_ERROR = None

if Document is None:  # pragma: no cover - skip module si docx indisponible
    pytest.skip(
        f"python-docx (package 'docx') requis pour tester les exports meeting report: {_DOCX_IMPORT_ERROR}",
        allow_module_level=True,
    )

_MINIMAL_PDF_BYTES = (
    b"%PDF-1.4\n"
    b"1 0 obj<<>>\nendobj\n"
    b"xref\n0 2\n0000000000 65535 f \n0000000009 00000 n \n"
    b"trailer<< /Size 2 >>\nstartxref\n44\n%%EOF\n"
)


@pytest.fixture
def fake_pandoc_conversion(monkeypatch):
    """
    Remplace la conversion pypandoc par un rendu local minimal pour les tests.
    """

    def _fake_convert(markdown_text: str, to: str, out_path):
        path = Path(out_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        if to == "docx":
            doc = Document()
            for line in markdown_text.splitlines():
                if line.startswith("### "):
                    doc.add_heading(line[4:].strip(), level=2)
                elif line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph("")
            doc.save(path)
        elif to == "pdf":
            path.write_bytes(_MINIMAL_PDF_BYTES)
        else:
            path.write_text(markdown_text, encoding="utf-8")

    monkeypatch.setattr(meeting_report_mod, "_convert_markdown_with_pandoc", _fake_convert)
    return _fake_convert


@pytest.fixture
def temp_project(tmp_path: Path):
    """
    Fabrique un mini "projet" temporaire avec :
      - un transcript anonymisé,
      - un mapping d'anonymisation,
      - un fichier de prompts JSON.
    Retourne un dict avec les chemins utiles.
    """
    root = tmp_path
    (root / "config").mkdir(parents=True, exist_ok=True)
    outputs_txt_dir = root / "outputs" / "txt"
    outputs_txt_dir.mkdir(parents=True, exist_ok=True)

    # 1) Transcription anonymisée nettoyée (out_txt_anon_clean)
    anon_txt = outputs_txt_dir / "testfile_anon_clean.txt"
    anon_txt.write_text(
        "Bonjour Alice Dupont. Nous avons évoqué le dossier Orion Conseil et fixé une échéance au 12/11.\n",
        encoding="utf-8",
    )

    # 2) Mapping anonymisation
    mapping_json = outputs_txt_dir / "testfile_anon_mapping.json"
    mapping = {
        "entities": {
            "<PERSON_1>": {
                "label": "PERSON",
                "values": ["Delphine"],
                "variants": ["Delphine"],
                "canonical": "Delphine",
                "pseudonym": "Alice Dupont",
                "source": "stub",
            },
            "<ORGANIZATION_1>": {
                "label": "ORGANIZATION",
                "values": ["UDAF"],
                "variants": ["UDAF"],
                "canonical": "UDAF",
                "pseudonym": "Orion Conseil",
                "source": "stub",
            },
        },
        "summary": {"PERSON": 1, "ORGANIZATION": 1},
        "reverse_map": {"<PERSON_1>": "Delphine", "<ORGANIZATION_1>": "UDAF"},
        "pseudonym_map": {"<PERSON_1>": "Alice Dupont", "<ORGANIZATION_1>": "Orion Conseil"},
        "pseudonym_reverse_map": {"Alice Dupont": "Delphine", "Orion Conseil": "UDAF"},
        "placeholder_style": "pseudonym",
        "stats": {"total": 2, "by_type": {"PERSON": 1, "ORGANIZATION": 1}},
        "meta": {"model_id": "stub"},
    }
    mapping_json.write_text(json.dumps(mapping, ensure_ascii=False, indent=2), encoding="utf-8")

    # 3) Prompts JSON (clé meeting_analysis)
    prompts_json = root / "config" / "mistral_prompts.json"
    prompts_json.write_text(
        json.dumps({
            "meeting_analysis": {
                "model": "mistral-large-latest",
                "system": "Tu es un assistant d'analyse de réunions professionnelles.",
                "user_prefix": "TRANSCRIPTION :\n"
            }
        }, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    return {
        "root": root,
        "anon_txt": anon_txt,
        "mapping_json": mapping_json,
        "prompts_json": prompts_json,
        "outputs_txt_dir": outputs_txt_dir,
    }


def _fake_llm_output():
    return (
        "### RÉSUMÉ EXÉCUTIF\n"
        "Point sur le dossier Orion Conseil. Décision d'avancer l'échéance.\n\n"
        "### PARTICIPANTS\n"
        "| Pseudonyme | Catégorie | Indices contextuels | Genre |\n"
        "| --- | --- | --- | --- |\n"
        "| Alice Dupont | Avocat gérante | Rôle de gestion et décisions | Féminin |\n\n"
        "### SUJETS ABORDÉS\n"
        "- Dossier — résumé: échéance au 12/11 «Nous devons finaliser le dossier». \n\n"
        "### DÉCISIONS\n"
        "- Avancer l’échéance au 12/11 — prise par Alice Dupont\n\n"
        "### ACTIONS\n"
        "- Envoyer les pièces — Responsable: Alice Dupont — Échéance: 12/11\n\n"
        "### PROCHAINES ÉTAPES\n"
        "- Réunion de suivi — Responsable: Alice Dupont — Délai: 1 semaine\n"
    )



def test_generate_meeting_report_end_to_end_no_network(
    temp_project, monkeypatch, fake_pandoc_conversion
):
    """
    Test end-to-end du rapport :
      - Mock de l'appel LLM (aucun réseau)
      - Mock de la désanonymisation via deanonymize_text()
      - Vérification des 3 sorties (txt anonymisé, txt désanonymisé, docx)
    """

    # === 1) Mock LLM : intercepter chat_complete et renvoyer une réponse prédictible ===
    def fake_chat_complete(model: str, system: str, user_text: str) -> str:
        # On pourrait vérifier ici que user_text contient la transcription anonymisée
        assert "TRANSCRIPTION" in user_text  # vient de user_prefix du prompt
        assert "Alice Dupont" in user_text
        assert "Pseudonymes détectés" in user_text  # rappel ajouté au prompt
        assert "Delphine" not in user_text
        return _fake_llm_output()

    # Patch ce que meeting_report.py utilise réellement :
    monkeypatch.setattr(meeting_report_mod.mistral_client, "chat_complete", fake_chat_complete)
    monkeypatch.setattr(meeting_report_mod, "_check_mistral_access", lambda timeout=5.0: (True, ""))

    # === 2) Mock désanonymisation : deanonymize_text ===
    # On remplace la fonction par une version simple qui applique le mapping fourni.
    def fake_deanon(text: str, mapping: dict, restore: str = "canonical") -> str:

        lookup = {}

        pseudo_reverse = mapping.get("pseudonym_reverse_map", {})
        if isinstance(pseudo_reverse, dict):
            lookup.update({k: v for k, v in pseudo_reverse.items() if k and v})

        reverse_map = mapping.get("reverse_map", {})
        if isinstance(reverse_map, dict):
            lookup.update({k: v for k, v in reverse_map.items() if k and v})

        entities = mapping.get("entities", {})
        if isinstance(entities, dict):
            for info in entities.values():
                pseudonym = info.get("pseudonym")
                canonical = info.get("canonical")
                if pseudonym and canonical and pseudonym not in lookup:
                    lookup[pseudonym] = canonical
        else:
            for entity in entities:
                tag = entity.get("tag")
                canonical = entity.get("canonical")
                if tag and canonical and tag not in lookup:
                    lookup[tag] = canonical

        out = text
        for key, real in lookup.items():
            out = out.replace(key, real)
        return out

    monkeypatch.setattr(meeting_report_mod, "deanonymize_text", fake_deanon)

    # === 3) Exécution ===
    root = temp_project["root"]
    anon_txt = temp_project["anon_txt"]
    mapping_json = temp_project["mapping_json"]
    prompts_json = temp_project["prompts_json"]
    outputs_txt_dir = temp_project["outputs_txt_dir"]

    result = generate_meeting_report(
        anonymized_txt_path=anon_txt,
        mapping_json_path=mapping_json,
        prompts_json_path=prompts_json,
        prompt_key="meeting_analysis",
        out_dir=outputs_txt_dir,       # on sort dans outputs/txt
        run_id="testfile",
    )

    # === 4) Vérifications des chemins retournés ===
    assert "report_anonymized_txt" in result
    assert "report_txt" in result
    assert "report_docx" in result
    assert "report_pdf" in result
    assert "report_markdown" in result
    assert result["report_status"] == "generated"
    assert result["report_reason"] == ""

    p_anon = Path(result["report_anonymized_txt"])
    p_txt = Path(result["report_txt"])
    p_docx = Path(result["report_docx"])
    p_pdf = Path(result["report_pdf"])
    p_md = Path(result["report_markdown"])

    assert p_anon.exists(), "Le .txt anonymisé du rapport doit exister"
    assert p_txt.exists(), "Le .txt désanonymisé du rapport doit exister"
    assert p_docx.exists(), "Le .docx du rapport doit exister"
    assert p_pdf.exists(), "Le .pdf du rapport doit exister"
    assert p_md.exists(), "Le .md du rapport doit exister"

    # === 5) Contenu des TXT ===
    anon_text = p_anon.read_text(encoding="utf-8")
    deanon_text = p_txt.read_text(encoding="utf-8")

    # La version anonymisée contient uniquement les pseudonymes
    assert "Alice Dupont" in anon_text and "Orion Conseil" in anon_text
    assert "Delphine" not in anon_text and "UDAF" not in anon_text
    assert "| Alice Dupont |" in anon_text

    # La version désanonymisée remplace bien par les vrais noms
    assert "Delphine" in deanon_text
    assert "UDAF" in deanon_text
    assert "Alice Dupont" not in deanon_text
    assert "Orion Conseil" not in deanon_text
    assert "| Delphine |" in deanon_text

    # Les sections ### sont bien présentes (utile pour le formatage docx)
    for section in [
        "### RÉSUMÉ EXÉCUTIF",
        "### PARTICIPANTS",
        "### SUJETS ABORDÉS",
        "### DÉCISIONS",
        "### ACTIONS",
        "### PROCHAINES ÉTAPES",
    ]:
        assert section in deanon_text

    # === 6) Lecture et checks rapides du DOCX ===
    doc = Document(str(p_docx))
    full_doc_text = "\n".join([p.text for p in doc.paragraphs])

    # Doit contenir les mêmes informations clé (désanonymisées)
    assert "Delphine" in full_doc_text
    assert "UDAF" in full_doc_text
    # Et au moins un titre de section transformé en heading
    assert any("RÉSUMÉ EXÉCUTIF" in p.text for p in doc.paragraphs)


def test_generate_meeting_report_skipped_when_mistral_unavailable(
    temp_project, monkeypatch
):
    """
    Vérifie que la génération est proprement court-circuitée si l'API Mistral est inaccessible.
    """

    monkeypatch.setattr(meeting_report_mod, "_check_mistral_access", lambda timeout=5.0: (False, "API indisponible"))

    root = temp_project["root"]
    anon_txt = temp_project["anon_txt"]
    mapping_json = temp_project["mapping_json"]
    prompts_json = temp_project["prompts_json"]

    result = generate_meeting_report(
        anonymized_txt_path=anon_txt,
        mapping_json_path=mapping_json,
        prompts_json_path=prompts_json,
        prompt_key="meeting_analysis",
        out_dir=root / "outputs" / "txt",
        run_id="testfile",
    )

    assert result["report_status"] == "skipped"
    assert "indisponible" in result["report_reason"]
    for key in [
        "report_anonymized_txt",
        "report_txt",
        "report_markdown",
        "report_docx",
        "report_pdf",
    ]:
        assert result[key] is None
