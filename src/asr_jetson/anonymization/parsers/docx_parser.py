"""DOCX parser extracting body text and selected metadata fields."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from asr_jetson.anonymization.core.models import ParsedDocument, Span


class DocxParser:
    def parse(self, input_path: Path, document_id: str) -> ParsedDocument:
        doc = Document(str(input_path))
        chunks: list[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                chunks.append(paragraph.text)
        core = doc.core_properties
        for attr in ("author", "title", "subject", "comments", "category"):
            value = getattr(core, attr, None)
            if value:
                chunks.append(str(value))

        text = "\n".join(chunks)
        span = Span(
            span_id=f"{document_id}:0",
            document_id=document_id,
            start=0,
            end=len(text),
            anchor_type="docx_xpath",
            anchor_ref="/word/document.xml",
        )
        return ParsedDocument(document_id=document_id, format="docx", text=text, spans=[span])
