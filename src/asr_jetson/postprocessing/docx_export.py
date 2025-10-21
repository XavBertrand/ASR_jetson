from __future__ import annotations

from typing import List, Optional, Sequence, Tuple
from xml.sax.saxutils import escape
from zipfile import ZipFile, ZIP_DEFLATED
import datetime

try:  # pragma: no cover - exercised via integration tests
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore
    _HAS_PYTHON_DOCX = True
except Exception:  # pragma: no cover - platform without python-docx
    Document = None  # type: ignore
    Pt = None  # type: ignore
    _HAS_PYTHON_DOCX = False


def _build_blocks(text: str, title: Optional[str]) -> List[Tuple[str, str, int]]:
    """
    Convert the markdown-flavoured text into logical blocks to render in DOCX.
    Blocks are tuples of (kind, content, heading_level).
    """
    blocks: List[Tuple[str, str, int]] = []
    buffer: List[str] = []

    if title:
        blocks.append(("heading", title, 1))

    def flush_paragraph() -> None:
        if buffer:
            blocks.append(("paragraph", "\n".join(buffer), 0))
            buffer.clear()

    for line in text.splitlines():
        if line.startswith("### "):
            flush_paragraph()
            blocks.append(("heading", line[4:].strip(), 2))
        else:
            buffer.append(line)

    flush_paragraph()
    if not blocks:
        blocks.append(("paragraph", "", 0))
    return blocks


def _write_with_python_docx(blocks: Sequence[Tuple[str, str, int]], out_path: str) -> None:
    doc = Document()  # type: ignore[call-arg]
    for kind, content, level in blocks:
        if kind == "heading":
            doc.add_heading(content, level=level if level else 1)
        else:
            doc.add_paragraph(content)

    # Normalise default text styling for better readability.
    style = doc.styles["Normal"]  # type: ignore[index]
    if Pt is not None:
        style.font.size = Pt(11)  # type: ignore[attr-defined]
    doc.save(out_path)


def _runs_xml(text: str) -> str:
    parts = text.split("\n")
    xml_parts = []
    for idx, part in enumerate(parts):
        xml_parts.append(f'<w:r><w:t xml:space="preserve">{escape(part)}</w:t></w:r>')
        if idx < len(parts) - 1:
            xml_parts.append("<w:br/>")
    if not xml_parts:
        xml_parts.append('<w:r><w:t/></w:r>')
    return "".join(xml_parts)


def _fallback_document_xml(blocks: Sequence[Tuple[str, str, int]]) -> str:
    body = []
    for kind, content, level in blocks:
        if kind == "heading":
            style = "Heading1" if level <= 1 else "Heading2"
            body.append(
                f'<w:p><w:pPr><w:pStyle w:val="{style}"/></w:pPr>{_runs_xml(content)}</w:p>'
            )
        else:
            body.append(f"<w:p>{_runs_xml(content)}</w:p>")
    body.append(
        "<w:sectPr>"
        "<w:pgSz w:w=\"12240\" w:h=\"15840\"/>"
        "<w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\"/>"
        "</w:sectPr>"
    )
    body_xml = "".join(body)
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
        'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
        'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
        'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" mc:Ignorable="w14 wp14">'
        f"<w:body>{body_xml}</w:body></w:document>"
    )


_CONTENT_TYPES_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
    <Default Extension="xml" ContentType="application/xml"/>
    <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
    <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
    <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
    <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>
"""

_RELS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
    <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
    <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""

_APP_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
            xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
    <Application>ASR Agent</Application>
    <DocSecurity>0</DocSecurity>
    <ScaleCrop>false</ScaleCrop>
    <HeadingPairs>
        <vt:vector size="2" baseType="variant">
            <vt:variant>
                <vt:lpstr>Paragraphs</vt:lpstr>
            </vt:variant>
            <vt:variant>
                <vt:i4>1</vt:i4>
            </vt:variant>
        </vt:vector>
    </HeadingPairs>
    <TitlesOfParts>
        <vt:vector size="1" baseType="lpstr">
            <vt:lpstr>Document</vt:lpstr>
        </vt:vector>
    </TitlesOfParts>
    <Company></Company>
    <LinksUpToDate>false</LinksUpToDate>
    <SharedDoc>false</SharedDoc>
    <HyperlinksChanged>false</HyperlinksChanged>
    <AppVersion>16.0000</AppVersion>
</Properties>
"""


def _iso_datetime() -> str:
    return datetime.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"


def _core_xml() -> str:
    now = _iso_datetime()
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
                   xmlns:dc="http://purl.org/dc/elements/1.1/"
                   xmlns:dcterms="http://purl.org/dc/terms/"
                   xmlns:dcmitype="http://purl.org/dc/dcmitype/"
                   xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
    <dc:title>Document</dc:title>
    <dc:creator>ASR Agent</dc:creator>
    <cp:lastModifiedBy>ASR Agent</cp:lastModifiedBy>
    <dcterms:created xsi:type="dcterms:W3CDTF">{now}</dcterms:created>
    <dcterms:modified xsi:type="dcterms:W3CDTF">{now}</dcterms:modified>
</cp:coreProperties>
"""


_STYLES_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
        <w:name w:val="Normal"/>
        <w:qFormat/>
    </w:style>
    <w:style w:type="paragraph" w:styleId="Heading1">
        <w:name w:val="heading 1"/>
        <w:basedOn w:val="Normal"/>
        <w:next w:val="Normal"/>
        <w:uiPriority w:val="9"/>
        <w:qFormat/>
    </w:style>
    <w:style w:type="paragraph" w:styleId="Heading2">
        <w:name w:val="heading 2"/>
        <w:basedOn w:val="Normal"/>
        <w:next w:val="Normal"/>
        <w:uiPriority w:val="9"/>
        <w:qFormat/>
    </w:style>
</w:styles>
"""


def _write_fallback_docx(blocks: Sequence[Tuple[str, str, int]], out_path: str) -> None:
    document_xml = _fallback_document_xml(blocks)
    with ZipFile(out_path, "w", ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", _CONTENT_TYPES_XML)
        zf.writestr("_rels/.rels", _RELS_XML)
        zf.writestr("docProps/app.xml", _APP_XML)
        zf.writestr("docProps/core.xml", _core_xml())
        zf.writestr("word/styles.xml", _STYLES_XML)
        zf.writestr("word/document.xml", document_xml)


def save_docx_from_markdown_sections(text: str, out_path: str, title: Optional[str] = None) -> str:
    """
    Create a DOCX file summarising the provided markdown-like content.
    Falls back to a minimal implementation when python-docx is unavailable.
    """
    blocks = _build_blocks(text, title)
    if _HAS_PYTHON_DOCX:
        _write_with_python_docx(blocks, out_path)
    else:
        _write_fallback_docx(blocks, out_path)
    return out_path
