"""
Module de génération de compte-rendu structuré avec **TensorRT-LLM** ou **Ollama**.
- Par défaut, utilise Ollama (moins de frictions d'installation).
- Optionnellement, peut utiliser TensorRT-LLM si disponible.

Usage exemples
--------------
# 1) Backend Ollama (par défaut) + pull auto du modèle s'il manque
python generate_report.py --input input.txt --output rapport.docx \
  --backend ollama --ollama-model qwen2.5:7b-instruct

# 2) Backend TensorRT-LLM (nécessite engines .engine déjà compilés)
python generate_report.py --input input.txt --output rapport.docx \
  --backend trtllm --model Athroniaeth/mistral-7b-v0.2-trtllm-int4

Notes
-----
- Ollama doit être en cours d'exécution en local (port 11434 par défaut).
  Installation: https://ollama.com/download
  Lancement: `ollama serve` (souvent auto-démarré)
"""

from pathlib import Path
from typing import Dict, List, Optional, Tuple
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import subprocess
import shutil
import time

# huggingface_hub n'est requis que si vous utilisez --backend trtllm et un repo HF
try:
    from huggingface_hub import snapshot_download  # type: ignore
except Exception:
    snapshot_download = None


def _find_engine_dir(root: Path) -> Optional[Path]:
    candidates = []
    for eng in root.rglob("*.engine"):
        candidates.append(eng.parent)
    for cfg in root.rglob("config.json"):
        if list(cfg.parent.glob("*.engine")):
            candidates.append(cfg.parent)
    if not candidates:
        return None
    candidates = sorted(candidates, key=lambda d: len(list(d.glob("*.engine"))), reverse=True)
    return candidates[0]


def _find_tokenizer_dir(root: Path) -> Path:
    for name in ("tokenizer.json", "tokenizer.model"):
        hit = next(root.rglob(name), None)
        if hit:
            return hit.parent
    return root


def resolve_model_sources(model_arg: str, cache_dir: Path = Path.home() / ".cache" / "trtllm_models") -> Tuple[
    Path, Optional[Path]]:
    path = Path(model_arg)
    if path.exists():
        root = path
    else:
        if snapshot_download is None:
            raise RuntimeError(
                "huggingface_hub n'est pas installé, requis pour télécharger un repo HF pour TensorRT-LLM.")
        target = cache_dir / model_arg.replace("/", "__")
        target.parent.mkdir(parents=True, exist_ok=True)
        root = Path(snapshot_download(repo_id=model_arg, local_dir=str(target)))
    tokenizer_dir = _find_tokenizer_dir(root)
    engine_dir = _find_engine_dir(root)
    return tokenizer_dir, engine_dir


# ============================================================================
# CONFIGURATION GÉNÉRALE (paramètres communs de génération)
# ============================================================================

class GenConfig:
    # Paramètres de génération par défaut (adaptés aux CR pros)
    TEMPERATURE = 0.3
    TOP_P = 0.85
    TOP_K = 40
    REPETITION_PENALTY = 1.15
    MAX_INPUT_LENGTH = 6144
    MAX_OUTPUT_LENGTH = 2048


# ============================================================================
# BACKEND: TensorRT-LLM
# ============================================================================

class TRTLLMConfig(GenConfig):
    MODEL_PATH = ""  # dossier tokenizer ou modèle
    ENGINE_DIR = ""  # dossier avec *.engine


def initialize_tensorrt_llm(config: TRTLLMConfig):
    try:
        import tensorrt_llm  # noqa: F401
        from tensorrt_llm.runtime import ModelRunner

        if not config.ENGINE_DIR or not Path(config.ENGINE_DIR).exists():
            raise RuntimeError(
                "ENGINE_DIR introuvable. Si vous avez passé un repo HF, "
                "assurez-vous qu'il contient des engines (*.engine). "
                "Sinon, compilez-les avec TensorRT-LLM."
            )
        runner = ModelRunner.from_dir(engine_dir=config.ENGINE_DIR, rank=0, debug_mode=False)
        return runner
    except ImportError:
        raise ImportError("TensorRT-LLM n'est pas installé. `pip install tensorrt_llm`")
    except Exception as e:
        raise RuntimeError(f"Erreur lors du chargement TensorRT-LLM: {e}")


def generate_with_tensorrt(runner, prompt: str, config: TRTLLMConfig, max_output_tokens: Optional[int] = None) -> str:
    max_out = max_output_tokens or config.MAX_OUTPUT_LENGTH
    from transformers import AutoTokenizer
    tokenizer = AutoTokenizer.from_pretrained(config.MODEL_PATH, use_fast=True, trust_remote_code=True)
    input_ids = tokenizer.encode(prompt, return_tensors="pt")
    outputs = runner.generate(
        input_ids,
        max_new_tokens=max_out,
        temperature=config.TEMPERATURE,
        top_p=config.TOP_P,
        top_k=config.TOP_K,
        repetition_penalty=config.REPETITION_PENALTY,
        end_id=tokenizer.eos_token_id,
        pad_id=tokenizer.pad_token_id,
    )
    output_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
    if prompt in output_text:
        output_text = output_text.replace(prompt, "").strip()
    return output_text


# ============================================================================
# BACKEND: Ollama
# ============================================================================

def _ollama_is_available() -> bool:
    # On cherche le binaire ollama (pour pull) et on suppose le daemon lancé
    return shutil.which("ollama") is not None


def ensure_ollama_model(model_name: str, timeout_s: int = 1800):
    """
    Fait un `ollama pull <model>` si le modèle n'est pas présent localement.
    On teste d'abord /api/tags puis, si absent, on pull via subprocess.
    """
    import requests

    try:
        resp = requests.get("http://localhost:11434/api/tags", timeout=5)
        if resp.ok:
            tags = resp.json().get("models", [])
            have = any(m.get("name") == model_name for m in tags)
            if have:
                return
    except Exception:
        # Si l'API n'est pas joignable, on tente quand même un pull (le daemon démarrera peut-être)
        pass

    if not _ollama_is_available():
        raise RuntimeError(
            "Ollama n'est pas installé ou non trouvé dans le PATH. Installez-le depuis https://ollama.com/.")

    print(f"   → Pull Ollama du modèle '{model_name}' (si nécessaire)...")
    try:
        proc = subprocess.Popen(["ollama", "pull", model_name], stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                                text=True)
    except FileNotFoundError:
        raise RuntimeError("Impossible d'exécuter 'ollama'. Vérifiez votre installation.")

    start = time.time()
    # Stream output pour feedback, et timeout de sécurité
    while True:
        if proc.poll() is not None:
            break
        if proc.stdout:
            line = proc.stdout.readline()
            if line:
                print("      ", line.rstrip())
        if time.time() - start > timeout_s:
            proc.kill()
            raise TimeoutError("Le pull du modèle Ollama a pris trop longtemps.")
    if proc.returncode != 0:
        raise RuntimeError("Échec du `ollama pull`. Consultez les logs ci-dessus.")


def generate_with_ollama(prompt: str, model_name: str, gen_cfg: GenConfig,
                         max_output_tokens: Optional[int] = None) -> str:
    """
    Utilise l'API locale d'Ollama (http://localhost:11434/api/generate)
    """
    import requests

    body = {
        "model": model_name,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.2,
            "top_p": 0.8,
            "top_k": 20,
            "num_predict": max_output_tokens or 2000,

            # anti-répétition
            "repeat_penalty": 1.4,
            "repeat_last_n": 512,
            "presence_penalty": 0.6,
            "frequency_penalty": 0.6,

            # stabilité
            "mirostat": 2,
            "mirostat_tau": 5.0,
            "mirostat_eta": 0.1,

            # VRAM
            "num_ctx": args.ollama_num_ctx,
            "num_batch": 512,
            "gpu_layers": 99,
        },
        "keep_alive": args.ollama_keep_alive,  # ex: "30s"
    }
    try:
        r = requests.post("http://localhost:11434/api/generate", json=body, timeout=600)
    except requests.exceptions.ConnectionError:
        raise RuntimeError("Impossible de se connecter à Ollama (http://localhost:11434). Lancez `ollama serve`.")
    if not r.ok:
        raise RuntimeError(f"Ollama API error: {r.status_code} {r.text}")
    data = r.json()
    return data.get("response", "").strip()


# ============================================================================
# CHUNKING + PROMPTS (identiques à la version TensorRT-LLM)
# ============================================================================

def smart_chunk_transcript(text: str, max_chars: int = 12000) -> List[str]:
    paragraphs = re.split(r'\n\n+|\n(?=SPEAKER_\d+:|SPK\d+:)', text)
    chunks = []
    current_chunk = []
    current_length = 0
    for para in paragraphs:
        para_len = len(para)
        if para_len > max_chars:
            sentences = re.split(r'(?<=[.!?])\s+', para)
            for sentence in sentences:
                if current_length + len(sentence) > max_chars:
                    if current_chunk:
                        chunks.append('\n\n'.join(current_chunk))
                        current_chunk = []
                        current_length = 0
                current_chunk.append(sentence)
                current_length += len(sentence)
        elif current_length + para_len <= max_chars:
            current_chunk.append(para)
            current_length += para_len
        else:
            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))
            current_chunk = [para]
            current_length = para_len
    if current_chunk:
        chunks.append('\n\n'.join(current_chunk))
    return chunks


SYSTEM_PROMPT = """Tu es un assistant d'analyse de réunions professionnelles.
Ton objectif est d'extraire uniquement des faits et de produire un rapport structuré, sans interprétation ni reformulation inutile.

⚠️ NE JAMAIS INVENTER NI SUPPOSER.
⚠️ NE PAS UTILISER DE FORMULATIONS GÉNÉRIQUES ("globalement", "il semble", "probablement").
⚠️ UTILISER UNIQUEMENT LES INFORMATIONS EXPLICITEMENT PRÉSENTES DANS LA TRANSCRIPTION.

Format attendu :
### RÉSUMÉ EXÉCUTIF
(5 à 10 lignes factuelles, sans adjectifs inutiles)

### PARTICIPANTS
- Nom/Rôle si identifiable (sinon "Intervenant inconnu")

### SUJETS ABORDÉS
- Thème — résumé (3 lignes max)
- Citations exactes entre guillemets quand pertinentes

### DÉCISIONS
- Décision — contexte — qui l’a prise

### ACTIONS
- Action — Responsable — Échéance

### PROCHAINES ÉTAPES
- Étape — Responsable — Délai

Ne sors rien en dehors de cette structure."""



def build_report_prompt(transcript_chunk: str, chunk_index: int, total_chunks: int) -> str:
    if total_chunks == 1:
        prompt = f"""{SYSTEM_PROMPT}

Analyse cette transcription d'entretien et réponds aux questions suivantes :

1. QUI parle avec QUI ? (déduis leurs rôles d'après la conversation)

2. QUEL est le contexte ? (type d'entretien, objectif)

3. QUELS sont les 5-7 THÈMES principaux abordés ?
   Pour chaque thème :
   - Titre du sujet
   - Résumé factuel (2-3 phrases)
   - Citations importantes

4. Y a-t-il des DÉCISIONS prises ? Lesquelles ?

5. Y a-t-il des ACTIONS à faire ? Lesquelles et par qui ?

6. CHIFFRES et DATES mentionnés (salaires, pourcentages, échéances)

7. POINTS POSITIFS exprimés

8. POINTS NÉGATIFS ou difficultés évoquées

9. PROCHAINES ÉTAPES mentionnées

Réponds en français, de manière structurée et factuelle.

TRANSCRIPTION :
{transcript_chunk}

ANALYSE :"""
    else:
        # Pour les chunks multiples
        prompt = f"""{SYSTEM_PROMPT}

Partie {chunk_index + 1}/{total_chunks} d'un long entretien.

Extrais de cette partie :

1. Thèmes abordés (avec résumé)
2. Décisions prises
3. Actions identifiées
4. Chiffres/dates mentionnés
5. Points importants

TRANSCRIPTION (Partie {chunk_index + 1}/{total_chunks}) :
{transcript_chunk}

EXTRACTION :"""

    return prompt


def merge_chunk_reports(chunk_reports: List[str]) -> str:
    merge_prompt = f"""
    Tu reçois {len(chunk_reports)} comptes rendus partiels de la même réunion.

    Tâche :
    - Ne PAS reformuler, ne PAS inventer.
    - Fusionne les sections en suivant le format :
      RÉSUMÉ EXÉCUTIF / PARTICIPANTS / SUJETS ABORDÉS / DÉCISIONS / ACTIONS / PROCHAINES ÉTAPES.
    - Ne garde que les éléments factuels, supprime les doublons.

    Sortie finale strictement dans ce format, sans introduction ni commentaire :
    """
    for i, report in enumerate(chunk_reports, 1):
        merge_prompt += f"\n### PARTIE {i}\n\n{report}\n\n"
    merge_prompt += "\n## COMPTE-RENDU UNIFIÉ ET COMPLET\n"
    return merge_prompt


# ============================================================================
# DOCX
# ============================================================================

def create_professional_word_document(report_content: str, output_path: Path, metadata: Optional[Dict] = None):
    doc = Document()
    title = doc.add_heading('COMPTE-RENDU DE RÉUNION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if metadata:
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Light Grid Accent 1'
        fields = {
            'Date': metadata.get('date', datetime.now().strftime('%d/%m/%Y')),
            'Durée': metadata.get('duration', 'N/A'),
            'Lieu': metadata.get('location', 'N/A'),
            'Rédacteur': metadata.get('author', 'Système automatique')
        }
        for key, value in fields.items():
            row = info_table.add_row()
            row.cells[0].text = key
            row.cells[1].text = str(value)
            row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    lines = report_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            heading = doc.add_heading(line.replace('### ', ''), level=1)
            heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith('#### '):
            heading = doc.add_heading(line.replace('#### ', ''), level=2)
            heading.runs[0].font.color.rgb = RGBColor(51, 102, 153)
        elif line.startswith('**') and '**' in line[2:]:
            p = doc.add_paragraph()
            bold_end = line.index('**', 2)
            bold_text = line[2:bold_end]
            rest_text = line[bold_end + 2:].strip()
            run_bold = p.add_run(bold_text)
            run_bold.bold = True
            run_bold.font.size = Pt(11)
            if rest_text:
                run_rest = p.add_run(' ' + rest_text)
                run_rest.font.size = Pt(11)
        elif line.startswith('- ') or line.startswith('• '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
    doc.add_page_break()
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Document généré automatiquement le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✓ Compte-rendu Word créé: {output_path}")


# ============================================================================
# PIPELINE PRINCIPALE
# ============================================================================

def generate_meeting_report(
        input_txt: Path,
        output_docx: Path,
        backend: str = "ollama",
        trt_config: Optional[TRTLLMConfig] = None,
        ollama_model: str = "qwen2.5:7b-instruct",
        metadata: Optional[Dict] = None
):
    """
    Génère un compte-rendu Word structuré depuis une transcription.
    - backend: 'ollama' (défaut) ou 'trtllm'
    - ollama_model: nom de modèle Ollama (ex: 'qwen2.5:7b-instruct', 'llama3.2:3b-instruct', 'mistral:7b')
    """
    gen_cfg = trt_config if trt_config is not None else GenConfig()

    print("=" * 60)
    print(f"   GÉNÉRATION DE COMPTE-RENDU ({backend.upper()})")
    print("=" * 60)

    # 1. Lecture de la transcription
    print("\n[1/5] Lecture de la transcription...")
    transcript = input_txt.read_text(encoding='utf-8')
    print(f"   ✓ Transcription chargée: {len(transcript)} caractères")

    # 2. Découpage intelligent
    print("\n[2/5] Analyse et découpage du texte...")
    chunks = smart_chunk_transcript(transcript, max_chars=12000)
    print(f"   ✓ Texte découpé en {len(chunks)} chunk(s)")

    # 3. Initialisation backend
    runner = None
    if backend == "trtllm":
        if trt_config is None:
            raise ValueError("Pour backend 'trtllm', fournissez un TRTLLMConfig initialisé.")
        print("\n[3/5] Initialisation de TensorRT-LLM...")
        runner = initialize_tensorrt_llm(trt_config)
        print("   ✓ Modèle TensorRT-LLM prêt")
    elif backend == "ollama":
        print("\n[3/5] Vérification du modèle Ollama...")
        ensure_ollama_model(ollama_model)
        print(f"   ✓ Modèle '{ollama_model}' prêt via Ollama")
    else:
        raise ValueError("backend doit être 'ollama' ou 'trtllm'")

    # 4. Génération
    print("\n[4/5] Génération du compte-rendu...")
    if len(chunks) == 1:
        prompt = build_report_prompt(chunks[0], 0, 1)
        if backend == "ollama":
            report = generate_with_ollama(prompt, ollama_model, gen_cfg)
        else:
            report = generate_with_tensorrt(runner, prompt, trt_config)  # type: ignore[arg-type]
        print("   ✓ Compte-rendu généré")
    else:
        chunk_reports = []
        for i, chunk in enumerate(chunks):
            print(f"   → Traitement du chunk {i + 1}/{len(chunks)}...")
            prompt = build_report_prompt(chunk, i, len(chunks))
            if backend == "ollama":
                partial = generate_with_ollama(prompt, ollama_model, gen_cfg, max_output_tokens=1024)
            else:
                partial = generate_with_tensorrt(runner, prompt, trt_config,
                                                 max_output_tokens=1024)  # type: ignore[arg-type]
            chunk_reports.append(partial)
        print("   → Fusion des analyses partielles...")
        merge_prompt = merge_chunk_reports(chunk_reports)
        if backend == "ollama":
            report = generate_with_ollama(merge_prompt, ollama_model, gen_cfg)
        else:
            report = generate_with_tensorrt(runner, merge_prompt, trt_config)  # type: ignore[arg-type]
        print("   ✓ Compte-rendu unifié généré")

    # 5. DOCX
    print("\n[5/5] Création du document Word...")
    create_professional_word_document(report, output_docx, metadata)

    print("\n" + "=" * 60)
    print("   ✓ TERMINÉ !")
    print(f"   Compte-rendu disponible: {output_docx}")
    print("=" * 60 + "\n")


# ============================================================================
# CLI
# ============================================================================

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Génère un compte-rendu avec Ollama ou TensorRT-LLM")
    parser.add_argument("--input", type=Path,
                        default="/home/xavier/ASR_Agent/outputs/txt/entretien_sylvie_silero_pyannote_pyannote_h2oai_faster-whisper-large-v3-turbo_clean.txt",
                        help="Fichier .txt de transcription nettoyée")
    parser.add_argument("--output", type=Path,
                        default="/home/xavier/ASR_Agent/outputs/txt/entretien_sylvie_silero_pyannote_pyannote_h2oai_faster-whisper-large-v3-turbo_clean.docx",
                        help="Chemin de sortie .docx")
    parser.add_argument("--backend", type=str, choices=["ollama", "trtllm"], default="ollama",
                        help="Backend de génération (par défaut: ollama)")
    parser.add_argument("--ollama-num-ctx", type=int, default=4096)
    parser.add_argument("--ollama-num-batch", type=int, default=1)
    parser.add_argument("--ollama-gpu-layers", type=int, default=99)
    parser.add_argument("--ollama-keep-alive", type=str, default="15s")

    # Options Ollama
    parser.add_argument("--ollama-model", type=str, default="deepseek-r1",
                        help="Nom du modèle Ollama (ex: 'qwen2.5:7b-instruct-q3_K_S', 'llama3.2:3b-instruct-q4_K_M', 'mistral:7b-instruct-q4_K_M')")
    # Options TRT-LLM
    parser.add_argument("--model", type=str, default="",
                        help="(TRT-LLM) Chemin local OU repo_id HF (avec engines .engine)")
    args = parser.parse_args()

    backend = args.backend

    trt_cfg: Optional[TRTLLMConfig] = None
    if backend == "trtllm":
        trt_cfg = TRTLLMConfig()
        if not args.model:
            raise SystemExit("Pour --backend trtllm, --model est requis.")
        tokenizer_dir, engine_dir = resolve_model_sources(args.model)
        trt_cfg.MODEL_PATH = str(tokenizer_dir)
        if engine_dir is not None:
            trt_cfg.ENGINE_DIR = str(engine_dir)
        else:
            print("⚠️  Aucun engine TRT-LLM détecté. Compilez d'abord les engines puis relancez.")
            raise SystemExit(1)

    generate_meeting_report(
        input_txt=args.input,
        output_docx=args.output,
        backend=backend,
        trt_config=trt_cfg,
        ollama_model=args.ollama_model,
        metadata={
            'date': datetime.now().strftime('%d/%m/%Y'),
            'duration': 'N/C',
            'location': 'N/C',
            'author': f'IA - {backend.upper()}'
        }
    )
